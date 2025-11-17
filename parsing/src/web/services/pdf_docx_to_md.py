import os
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def docx_to_markdown(docx_path, output_path=None):
    """
    Конвертирует .docx файл в markdown формат.
    
    Args:
        docx_path: путь к .docx файлу
        output_path: путь для сохранения markdown файла (опционально)
    
    Returns:
        str: содержимое документа в markdown формате
    """
    doc = Document(docx_path)
    markdown_lines = []
    
    def get_paragraph_style_level(paragraph):
        """Определяет уровень заголовка по стилю параграфа"""
        style_name = paragraph.style.name.lower()
        if 'heading' in style_name:
            try:
                level = int(style_name.replace('heading', '').strip())
                return min(level, 6)  # Markdown поддерживает только 6 уровней
            except:
                return 1
        return None
    
    def process_run(run):
        """Обрабатывает run (фрагмент текста) с форматированием"""
        text = run.text
        if not text:
            return ""
        
        # Жирный текст
        if run.bold:
            text = f"**{text}**"
        
        # Курсив
        if run.italic:
            text = f"*{text}*"
        
        # Подчеркивание (в markdown нет нативного подчеркивания, используем <u>)
        if run.underline:
            text = f"<u>{text}</u>"
        
        # Зачеркнутый текст
        if run.font.strike or (hasattr(run, '_element') and run._element.get(qn('w:strike'))):
            text = f"~~{text}~~"
        
        return text
    
    def process_paragraph(paragraph):
        """Обрабатывает параграф"""
        # Пропускаем пустые параграфы
        if not paragraph.text.strip() and not paragraph.runs:
            return ""
        
        # Проверяем, является ли параграф заголовком
        heading_level = get_paragraph_style_level(paragraph)
        if heading_level:
            text = paragraph.text.strip()
            if text:
                markdown_lines.append(f"{'#' * heading_level} {text}\n")
            return
        
        # Обрабатываем текст параграфа с форматированием
        paragraph_text = ""
        for run in paragraph.runs:
            paragraph_text += process_run(run)
        
        # Проверяем, является ли параграф частью списка
        if paragraph.style.name.startswith('List'):
            # Определяем тип списка (нумерованный или маркированный)
            if 'number' in paragraph.style.name.lower() or 'numbered' in paragraph.style.name.lower():
                # Для нумерованных списков можно использовать автоматическую нумерацию
                markdown_lines.append(f"1. {paragraph_text.strip()}\n")
            else:
                markdown_lines.append(f"- {paragraph_text.strip()}\n")
        else:
            if paragraph_text.strip():
                markdown_lines.append(f"{paragraph_text.strip()}\n")
    
    def process_table(table):
        """Обрабатывает таблицу"""
        if not table.rows:
            return
        
        # Заголовок таблицы (первая строка)
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        markdown_lines.append("| " + " | ".join(headers) + " |\n")
        markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |\n")
        
        # Остальные строки
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # Дополняем ячейки, если их меньше, чем заголовков
            while len(cells) < len(headers):
                cells.append("")
            markdown_lines.append("| " + " | ".join(cells) + " |\n")
        
        markdown_lines.append("\n")
    
    # Обрабатываем все элементы документа
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Параграф
            paragraph = None
            for p in doc.paragraphs:
                if p._element == element:
                    paragraph = p
                    break
            if paragraph:
                process_paragraph(paragraph)
        elif element.tag.endswith('tbl'):  # Таблица
            table = None
            for t in doc.tables:
                if t._element == element:
                    table = t
                    break
            if table:
                process_table(table)
    
    # Объединяем все строки
    markdown_content = "".join(markdown_lines)
    
    # Сохраняем в файл, если указан путь
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Markdown файл сохранен: {output_path}")
    
    return markdown_content

# Пример использования
if __name__ == "__main__":
    # Тестируем на одном из файлов
    test_file = "/Users/egorbykov/Desktop/Работа/2025/hackatons/wink/data/1 кейс Сценарии c Таблицами/Фишер (с таблицей)/1/ФИШЕР 1 сери __16.05.docx"
    if os.path.exists(test_file):
        markdown_text = docx_to_markdown(test_file)
        print("Первые 1000 символов markdown:")
        print(markdown_text[:1000])
    else:
        print(f"Файл не найден: {test_file}")


import os
import re
from docx import Document
from docx.oxml.ns import qn

def docx_to_markdown_advanced(docx_path, output_path=None):
    """
    Улучшенная версия парсера .docx в markdown.
    Поддерживает:
    - Заголовки всех уровней (H1-H6)
    - Форматирование текста (жирный, курсив, подчеркивание, зачеркнутый)
    - Таблицы
    - Гиперссылки
    - Правильный порядок элементов (параграфы и таблицы в нужной последовательности)

    >>> Изменено: Логика автоматической нумерации списков (w:listText) отключена
    для сохранения нестандартных текстовых префиксов (например, "4-1. ИНТ..."),
    которые Word может скрывать как автоматическую нумерацию.

    Для обработки списков используется упрощенная эвристика.
    """
    doc = Document(docx_path)
    markdown_lines = []

    # Создаем словари для быстрого доступа к параграфам и таблицам по элементам
    paragraphs_dict = {p._element: p for p in doc.paragraphs}
    tables_dict = {t._element: t for t in doc.tables}

    def get_heading_level(paragraph):
        """Определяет уровень заголовка по стилю"""
        style_name = paragraph.style.name.lower()
        if 'heading' in style_name or 'заголовок' in style_name:
            # Извлекаем номер из названия стиля
            for word in style_name.split():
                if word.isdigit():
                    level = int(word)
                    return min(max(level, 1), 6)
            # Попытка определить по стандартным названиям
            if 'heading 1' in style_name or 'заголовок 1' in style_name:
                return 1
            elif 'heading 2' in style_name or 'заголовок 2' in style_name:
                return 2
            elif 'heading 3' in style_name or 'заголовок 3' in style_name:
                return 3
            elif 'heading 4' in style_name or 'заголовок 4' in style_name:
                return 4
            elif 'heading 5' in style_name or 'заголовок 5' in style_name:
                return 5
            elif 'heading 6' in style_name or 'заголовок 6' in style_name:
                return 6
        return None

    # --- Измененные функции списков (упрощены, чтобы не мешать нестандартной нумерации) ---
    def is_list_paragraph(paragraph):
        """Проверяет, является ли параграф элементом списка через XML"""
        p_element = paragraph._element
        numPr = p_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        # Включаем логику списков ТОЛЬКО для параграфов с явным списковым стилем
        # (это должно уменьшить вероятность конфликта со сценарной нумерацией)
        style_name = paragraph.style.name.lower()
        is_list_style = 'list' in style_name or 'список' in style_name or 'bullet' in style_name or 'маркер' in style_name
        return numPr is not None or is_list_style

    def get_list_info(paragraph):
        """Получает информацию о списке (тип и уровень) - упрощенная версия"""
        p_element = paragraph._element
        numPr = p_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')

        level = 0
        if numPr is not None:
            # Определяем уровень вложенности
            ilvl = numPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
            if ilvl is not None and ilvl.get(qn('w:val')) is not None:
                try:
                    level = int(ilvl.get(qn('w:val')))
                except:
                    pass

        # Определяем тип списка (нумерованный или маркированный) по стилю
        style_name = paragraph.style.name.lower()
        is_numbered = 'number' in style_name or 'нумерованный' in style_name or ('list paragraph' in style_name and 'bullet' not in style_name)

        if not is_numbered:
             # Эвристика: если numPr есть, но это не нумерованный стиль, считаем маркированным
            if numPr is not None and 'bullet' not in style_name and 'маркер' not in style_name:
                 is_numbered = True # По умолчанию, если нет маркера, но есть нумерация, это нумерованный

        # Возвращаем (is_numbered, level) только если это действительно элемент списка
        if is_list_paragraph(paragraph):
            return is_numbered, level
        return None, 0 # Это не элемент списка

    # --- Конец измененных функций списков ---

    def process_run(run):
        """Обрабатывает run (фрагмент текста) с форматированием"""
        text = run.text
        if not text:
            return ""

        # Обработка гиперссылок
        try:
            hyperlinks = run._element.xpath('.//w:hyperlink',
                namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if hyperlinks:
                for hyperlink in hyperlinks:
                    rId = hyperlink.get(qn('r:id'))
                    if rId and rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        url = rel.target_ref
                        # Извлекаем текст ссылки из элемента
                        link_texts = []
                        # NOTE: Изменение - ищем текст во всех run внутри w:hyperlink
                        for r_elem in hyperlink.xpath('.//w:r', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for t in r_elem.iter():
                                if t.text:
                                    link_texts.append(t.text)
                        link_text = ''.join(link_texts)
                        if link_text:
                            # Заменяем только первое вхождение, чтобы не сломать, если текст run совпадает с текстом ссылки
                            # Этот метод не идеален, но работает для большинства случаев
                            # Если run.text - это часть ссылки, он будет заменен
                            text = text.replace(link_text, f"[{link_text}]({url})", 1)
        except:
            pass

        # Жирный текст
        if run.bold:
            text = f"**{text}**"

        # Курсив
        if run.italic:
            text = f"*{text}*"

        # Подчеркивание
        if run.underline:
            text = f"<u>{text}</u>"

        # Зачеркнутый текст
        try:
            strike_elem = run._element.find(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}strike')
            if strike_elem is not None:
                text = f"~~{text}~~"
        except:
            pass

        return text

    def process_paragraph(paragraph):
        """Обрабатывает параграф"""
        # Собираем текст из всех runs с форматированием
        text_content = ""
        for run in paragraph.runs:
            text_content += process_run(run)

        text_content = text_content.strip()

        # Пропускаем пустые параграфы (кроме случаев, когда это важно для структуры)
        if not text_content and not paragraph.runs:
            return

        # Проверяем заголовки
        heading_level = get_heading_level(paragraph)
        if heading_level:
            if text_content:
                markdown_lines.append(f"{'#' * heading_level} {text_content}\n\n")
            return

        # Проверяем, является ли параграф элементом списка
        is_numbered, list_level = get_list_info(paragraph)

        if is_numbered is not None:
            # Это элемент списка

            # --- КЛЮЧЕВОЕ ИЗМЕНЕНИЕ ДЛЯ РЕШЕНИЯ ПРОБЛЕМЫ ---
            # Если Word скрыл автоматическую нумерацию ("4-1."),
            # мы обрабатываем этот параграф как обычный, чтобы сохранить текстовый префикс,
            # но только если параграф НЕ пуст.

            # Эвристика: Если текст начинается с цифры, это может быть нестандартная нумерация.
            # Если текст начинается с цифры И НЕ содержит пробела после первой цифры (например, '4-1.ИНТ'),
            # или если это точно не заголовок, мы его пропускаем как список.

            # Внимание: Поскольку мы не можем получить скрытый "4-1." из docx.paragraph.text,
            # единственное, что мы можем сделать - это ПРЕДПОЛОЖИТЬ,
            # что если это нумерованный список, и в нем нет текста, то мы не добавляем "1. ",
            # но в вашем случае текст ЕСТЬ, просто без "4-1.".

            # Для вашего случая (4-1. ИНТ...):
            # 1. Если "4-1." является автоматическим listText, text_content будет "ИНТ. КОЛЛЕДЖ..."
            # 2. Логика ниже добавит "1. ", и вы получите "1. ИНТ. КОЛЛЕДЖ..."

            # Чтобы сохранить "ИНТ. КОЛЛЕДЖ..." БЕЗ "1. ":
            # Если мы хотим сохранить "4-1. ИНТ..." как обычный параграф, нужно пропустить логику списка.

            # Но если это настоящий список:
            indent = "  " * list_level

            if is_numbered:
                # Нумерованный список: используем 1., Markdown сам исправит
                # Мы не добавляем 1. , если текст уже начинается с цифры (для предотвращения дублирования)
                if not re.match(r'^\d+\.?\s*|\w+\.\s*', text_content): # Если текст НЕ начинается с нумерации
                    markdown_lines.append(f"{indent}1. {text_content}\n")
                else:
                    # Если текст уже содержит нумерацию (напр. "4-1. ИНТ..."), обрабатываем как обычный текст
                    markdown_lines.append(f"{text_content}\n")
            else:
                # Маркированный список
                markdown_lines.append(f"{indent}- {text_content}\n")
        else:
            # Обычный параграф (включая параграфы, которые были списками с нестандартным префиксом)
            if text_content:
                markdown_lines.append(f"{text_content}\n\n")
            elif not markdown_lines or markdown_lines[-1].strip() != "":
                 # Добавляем пустую строку только если предыдущая строка не была пустой
                 markdown_lines.append("\n")

    def process_table(table):
        """Обрабатывает таблицу"""
        if not table.rows:
            return

        markdown_lines.append("\n")  # Отступ перед таблицей

        # Обрабатываем каждую строку
        for i, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                # Обрабатываем содержимое ячейки (могут быть параграфы)
                cell_texts = []
                for para in cell.paragraphs:
                    para_text = ""
                    for run in para.runs:
                        para_text += process_run(run)
                    if para_text.strip():
                        cell_texts.append(para_text.strip())
                cell_text = " ".join(cell_texts) if cell_texts else " "
                cell_text = cell_text.replace('\n', ' ').strip()
                cells.append(cell_text if cell_text else " ")

            if cells:
                markdown_lines.append("| " + " | ".join(cells) + " |\n")

                # Добавляем разделитель после первой строки
                if i == 0:
                    markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |\n")

        markdown_lines.append("\n")  # Отступ после таблицы

    # Обрабатываем элементы в правильном порядке (как они идут в документе)
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Параграф
            paragraph = paragraphs_dict.get(element)
            if paragraph:
                process_paragraph(paragraph)
        elif element.tag.endswith('tbl'):  # Таблица
            table = tables_dict.get(element)
            if table:
                process_table(table)

    # Объединяем результат
    markdown_content = "".join(markdown_lines)

    # Очищаем лишние пустые строки (более 2 подряд)
    markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)

    # Убираем лишние пробелы в начале и конце
    markdown_content = markdown_content.strip() + '\n'

    # Сохраняем файл
    if output_path:
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        # print(f"✓ Markdown файл сохранен: {output_path}") # Закомментировано для чистоты вывода

    return markdown_content


import os
import re
import fitz  # PyMuPDF

# Парсинг PDF файлов в Markdown

import re
import fitz  # PyMuPDF

def pdf_to_markdown(pdf_path, output_path=None):
    """
    Конвертирует .pdf файл в markdown формат.
    
    Args:
        pdf_path: путь к .pdf файлу
        output_path: путь для сохранения markdown файла (опционально)
    
    Returns:
        str: содержимое документа в markdown формате
    """
    doc = fitz.open(pdf_path)
    markdown_lines = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Извлекаем текст со структурой
        blocks = page.get_text("dict")
        
        for block in blocks["blocks"]:
            if "lines" in block:  # Текстовый блок
                paragraph_text = ""
                is_bold = False
                is_italic = False
                
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        text = span["text"]
                        font_flags = span.get("flags", 0)
                        
                        # Определяем форматирование
                        is_bold_span = font_flags & 16  # Bold flag
                        is_italic_span = font_flags & 2  # Italic flag
                        
                        # Применяем форматирование
                        if is_bold_span:
                            text = f"**{text}**"
                        if is_italic_span:
                            text = f"*{text}*"
                        
                        line_text += text
                    
                    if line_text.strip():
                        paragraph_text += line_text + " "
                
                paragraph_text = paragraph_text.strip()
                
                # Проверяем, является ли текст заголовком (по размеру шрифта)
                if block["lines"]:
                    first_span = block["lines"][0]["spans"][0]
                    font_size = first_span.get("size", 12)
                    
                    # Эвристика: больший шрифт = заголовок
                    if font_size >= 16:
                        heading_level = 1
                    elif font_size >= 14:
                        heading_level = 2
                    elif font_size >= 12.5:
                        heading_level = 3
                    else:
                        heading_level = None
                    
                    if heading_level and paragraph_text:
                        markdown_lines.append(f"{'#' * heading_level} {paragraph_text}\n\n")
                    elif paragraph_text:
                        markdown_lines.append(f"{paragraph_text}\n\n")
            
            # Обработка изображений (опционально)
            elif "image" in block:
                # Можно добавить ссылки на изображения
                pass
        
        # Извлекаем таблицы
        tables = page.find_tables()
        for table in tables:
            try:
                table_data = table.extract()
                if table_data and len(table_data) > 0:
                    # Заголовок таблицы
                    headers = table_data[0]
                    markdown_lines.append("| " + " | ".join([str(h) if h else " " for h in headers]) + " |\n")
                    markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |\n")
                    
                    # Строки таблицы
                    for row in table_data[1:]:
                        row_cells = [str(cell) if cell else " " for cell in row]
                        # Дополняем строку, если не хватает ячеек
                        while len(row_cells) < len(headers):
                            row_cells.append(" ")
                        markdown_lines.append("| " + " | ".join(row_cells) + " |\n")
                    
                    markdown_lines.append("\n")
            except Exception as e:
                # Если не удалось извлечь таблицу, пропускаем
                pass
        
        # Добавляем разделитель страниц (опционально)
        if page_num < len(doc) - 1:
            markdown_lines.append("---\n\n")
    
    doc.close()
    
    # Объединяем все строки
    markdown_content = "".join(markdown_lines)
    
    # Очищаем лишние пустые строки
    markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
    markdown_content = markdown_content.strip() + '\n'
    
    # Сохраняем в файл, если указан путь
    if output_path:
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"✓ Markdown файл сохранен: {output_path}")
    
    return markdown_content


import fitz
import re
import os
import statistics

def pdf_to_markdown_advanced(pdf_path, output_path=None):
    """
    Улучшенная версия парсера PDF → Markdown:
    - Сохраняет структуру текста
    - Определяет заголовки на основе относительного размера шрифта
    - Поддерживает списки и форматирование
    - Устраняет ложные таблицы
    - Корректно формирует абзацы
    """
    doc = fitz.open(pdf_path)
    markdown_lines = []

    for page_num, page in enumerate(doc):
        text_dict = page.get_text("dict")
        all_font_sizes = []

        # Сначала собираем все размеры шрифтов (для относительной шкалы)
        for block in text_dict["blocks"]:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if "size" in span:
                        all_font_sizes.append(span["size"])

        avg_font_size = statistics.mean(all_font_sizes) if all_font_sizes else 12

        def is_likely_heading(text, font_size, is_bold):
            """Определяет, является ли текст заголовком (относительно средней высоты шрифта)"""
            if not text or len(text.strip()) < 2:
                return False, 0
            text_upper_ratio = sum(1 for c in text if c.isupper()) / len(text) if text else 0
            is_short = len(text.split()) <= 10

            if font_size > avg_font_size * 1.8:
                return True, 1
            elif font_size > avg_font_size * 1.5 or (is_bold and is_short):
                return True, 2
            elif is_bold and text_upper_ratio > 0.5:
                return True, 3
            return False, 0

        previous_y = None
        current_paragraph = []

        for block in text_dict["blocks"]:
            if "lines" not in block:
                continue

            for line in block["lines"]:
                line_text = ""
                line_y = line["bbox"][1]  # верхняя граница строки
                line_font_sizes = []
                line_bold_flags = []

                for span in line["spans"]:
                    text = span.get("text", "").strip()
                    if not text:
                        continue

                    font_size = span.get("size", avg_font_size)
                    font_flags = span.get("flags", 0)
                    is_bold = bool(font_flags & 16)
                    is_italic = bool(font_flags & 2)

                    line_font_sizes.append(font_size)
                    line_bold_flags.append(is_bold)

                    if is_bold:
                        text = f"**{text}**"
                    if is_italic:
                        text = f"*{text}*"

                    line_text += text + " "

                if not line_text.strip():
                    continue

                avg_line_font = statistics.mean(line_font_sizes) if line_font_sizes else avg_font_size
                is_line_bold = any(line_bold_flags)

                # Проверяем, не заголовок ли строка
                is_heading, level = is_likely_heading(line_text, avg_line_font, is_line_bold)
                if is_heading:
                    if current_paragraph:
                        markdown_lines.append(" ".join(current_paragraph).strip() + "\n\n")
                        current_paragraph = []
                    markdown_lines.append(f"{'#' * level} {line_text.strip()}\n\n")
                    continue

                # Определяем границу абзаца по вертикальному отступу
                if previous_y is not None and (line_y - previous_y) > avg_font_size * 1.5:
                    markdown_lines.append(" ".join(current_paragraph).strip() + "\n\n")
                    current_paragraph = []

                current_paragraph.append(line_text.strip())
                previous_y = line_y

        # Завершаем последний параграф
        if current_paragraph:
            markdown_lines.append(" ".join(current_paragraph).strip() + "\n\n")

        # Поиск таблиц (но фильтруем ложные)
        try:
            tables = page.find_tables()
            for table in tables:
                table_data = table.extract()
                if not table_data or len(table_data) < 2:
                    continue

                # Проверяем, что это действительно таблица
                text_join = " ".join(" ".join(str(c) for c in r) for r in table_data)
                if not any(ch in text_join for ch in [",", ";", " "]):
                    continue

                headers = [str(h).strip() if h else " " for h in table_data[0]]
                markdown_lines.append("\n| " + " | ".join(headers) + " |\n")
                markdown_lines.append("| " + " | ".join(["---"] * len(headers)) + " |\n")

                for row in table_data[1:]:
                    cells = [str(c).strip().replace("\n", " ") if c else " " for c in row]
                    while len(cells) < len(headers):
                        cells.append(" ")
                    markdown_lines.append("| " + " | ".join(cells) + " |\n")

                markdown_lines.append("\n")
        except Exception:
            pass

        # Добавляем разделитель страниц
        if page_num < len(doc) - 1:
            markdown_lines.append("---\n\n")

    doc.close()

    markdown_content = "\n".join(markdown_lines)
    markdown_content = re.sub(r"\n{3,}", "\n\n", markdown_content).strip() + "\n"

    if output_path:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        print(f"✓ Markdown файл сохранен: {output_path}")

    return markdown_content


def file_to_markdown(file_path, output_path=None):
    """
    Универсальная функция для конвертации .docx или .pdf в Markdown.

    Args:
        file_path: путь к исходному файлу (.docx или .pdf)
        output_path: путь для сохранения markdown файла (опционально)

    Returns:
        str: содержимое документа в markdown формате
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    
    if output_path is None:
        output_path = os.path.splitext(file_path)[0] + ".md"

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        # Используем docx_to_markdown_advanced

        markdown_text = docx_to_markdown_advanced(file_path, output_path)
    elif ext == ".pdf":
        # Используем pdf_to_markdown_advanced

        markdown_text = pdf_to_markdown_advanced(file_path, output_path)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}. Поддерживаются только .docx и .pdf")

    return markdown_text
